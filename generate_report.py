from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

# Load template
template_path = r"c:\Users\Programmer\Downloads\503363-programming-project-report-template (1) (1).docx"
doc = Document(template_path)

# Screenshots folder
screenshots_dir = r"C:\Users\Programmer\Desktop\Developments\Side Projects\chess-game\screenshots"

# Read code files
code_dir = r"C:\Users\Programmer\Desktop\Developments\Side Projects\chess-game"
with open(os.path.join(code_dir, "chessMain.py"), "r", encoding="utf-8") as f:
    chess_main_code = f.read()
with open(os.path.join(code_dir, "chessEngine.py"), "r", encoding="utf-8") as f:
    chess_engine_code = f.read()

# Fill date table (Table 0)
if len(doc.tables) > 0:
    date_table = doc.tables[0]
    if len(date_table.rows) > 1:
        date_table.rows[1].cells[0].text = datetime.now().strftime("%Y-%m-%d")
        date_table.rows[1].cells[1].text = datetime.now().strftime("%Y-%m-%d")

# Fill Test Design Table (Table 1)
if len(doc.tables) > 1:
    test_table = doc.tables[1]
    test_data = [
        ("1", "Login with valid user", "Username: test, Password: test", "Normal", "Login succeeds and settings window appears"),
        ("2", "Login with wrong password", "Username: test, Password: wrong", "Erroneous", "Error message \"Invalid username or password\""),
        ("3", "Register new user", "New username and password", "Normal", "Message \"Registered. You can log in now.\" and can then log in"),
        ("4", "Making a legal move", "Click e2 then e4 (White)", "Normal", "Pawn moves to e4; turn switches to Black"),
        ("5", "Move when time is zero", "Let one player's clock reach 0:00", "Boundary", "Game shows \"White/Black out of time - [other] wins\" and restart option")
    ]
    
    for i, row_data in enumerate(test_data):
        if i + 1 < len(test_table.rows):
            row = test_table.rows[i + 1]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text

# Fill Testing Results Table (Table 2)
if len(doc.tables) > 2:
    results_table = doc.tables[2]
    results_data = [
        ("1", "Login with valid user", "Login succeeds, settings appear", "Pass", "No change"),
        ("2", "Login with wrong password", "Error message shown", "Pass", "No change"),
        ("3", "Register new user", "Registration and login work", "Pass", "No change"),
        ("4", "Legal move (e2–e4)", "Pawn moves, turn changes", "Pass", "No change"),
        ("5", "Clock reaches zero", "Time-out message and winner shown", "Pass", "No change")
    ]
    
    for i, row_data in enumerate(results_data):
        if i + 1 < len(results_table.rows):
            row = results_table.rows[i + 1]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text

# Process paragraphs - find sections and add content
# We'll collect what to add and process in reverse order
additions = []  # List of (para_index, content_list) tuples

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Analysis - Success Criteria
    if text == "Success Criteria:":
        criteria = [
            "1. User authentication – The program allows users to log in or register; credentials are stored (e.g. in users.csv) and checked before starting the game.",
            "2. Correct chess rules – All pieces move according to standard rules (pawns, knights, bishops, rooks, queen, king), including castling, en passant, pawn promotion, and check/checkmate/stalemate detection.",
            "3. Playable interface – The game displays an 8×8 board with piece images, file/rank labels, move highlights for the selected piece, and valid-move indicators (e.g. dots).",
            "4. Time controls – Each side has a clock; time decreases on their turn, with optional increment per move. The game ends when one side runs out of time.",
            "5. Game end and restart – The program shows a clear message for checkmate, stalemate, or time-out, and allows the user to restart (e.g. Play Again / New Game or key R/N)."
        ]
        additions.append((i, criteria))
    
    # Design - Flow
    if "flow charts" in text.lower() and "broadly" in text.lower():
        flow_text = [
            "Flow (broad behaviour):",
            "",
            "1. Show login window (Tkinter) → user enters username/password or registers.",
            "2. If login OK → show settings window: choose side (White/Black), time control (e.g. 1+0, 3+2, 15+10).",
            "3. Start Pygame window: load board and piece images, set initial game state and clocks.",
            "4. Main loop: handle events (quit, mouse click, keys). On click: select square or make move if two squares selected; validate move using engine; update board and clocks; handle promotion pop-up if needed.",
            "5. Each frame: decrement current player's time; redraw board, pieces, highlights, move history, clocks; if game over, show message and restart menu.",
            "6. User can undo (e.g. U), restart (R/N), or quit."
        ]
        additions.append((i, flow_text))
    
    # Design - Pseudocode
    if "pseudocode" in text.lower() and "minimum" in text.lower():
        pseudocode = [
            "Pseudocode (move validation and making a move):",
            "",
            "FUNCTION getValidMoves():",
            "    moves = empty list",
            "    FOR each row r from 0 to 7:",
            "        FOR each column c from 0 to 7:",
            "            piece = board[r][c]",
            "            IF piece is empty OR piece colour is not current player's turn:",
            "                CONTINUE to next square",
            "            CALL appropriate move function for piece type (pawn, rook, knight, etc.) with (r, c, moves)",
            "    CALL addCastleMoves(moves)",
            "    legal = empty list",
            "    FOR each move m in moves:",
            "        IF m would capture a king:",
            "            CONTINUE",
            "        save current enPassant state",
            "        CALL makeMove(m)",
            "        IF current player's king is NOT in check after move:",
            "            ADD m to legal",
            "        CALL undoMove()",
            "    RETURN legal",
            "",
            "FUNCTION makeMove(move):",
            "    SET board[move.startRow][move.startCol] to empty",
            "    IF move is en passant:",
            "        REMOVE captured pawn from board",
            "    SET board[move.endRow][move.endCol] to move.pieceMoved",
            "    IF piece moved is king:",
            "        UPDATE king location for that colour",
            "    IF move is castle:",
            "        MOVE rook to correct square",
            "    IF move is pawn promotion:",
            "        SET promoted square to queen (or chosen piece later)",
            "    UPDATE en passant possibility if pawn moved two squares",
            "    UPDATE castle rights",
            "    APPEND move to moveLog",
            "    SWITCH turn to other player"
        ]
        additions.append((i, pseudocode))
    
    # Development - Code
    if "My program code:" in text:
        code_lines = chess_main_code.split('\n')
        code_lines2 = chess_engine_code.split('\n')
        code_content = [
            "chessMain.py - Main game: login, settings, Pygame board, move input, clocks, promotion, game over.",
            ""
        ] + code_lines + [
            "",
            "chessEngine.py - Game state, move generation, check/checkmate, castling, en passant.",
            ""
        ] + code_lines2
        additions.append((i, code_content))
    
    # Evaluation sections
    if "How successful was my program?" in text:
        success_text = [
            "The program meets the success criteria: users can log in and register, choose side and time control, and play a full game of chess on an 8×8 board with correct rules. All piece types move correctly, including castling and en passant, and check, checkmate and stalemate are detected. The interface shows move highlights and the last few moves, and both players have clocks that count down with optional increment. When the game ends (checkmate, stalemate or time-out), a clear message is shown and the user can restart. Testing showed that valid and invalid logins, registration, normal moves, and time-out behave as expected. One limitation is that the game is two-player on one machine (no network or AI opponent). Another is that promotion defaults to queen in the engine, but the GUI allows choosing queen, rook, bishop or knight. Overall, the program is successful for its intended purpose as a local two-player chess game with timing and standard rules."
        ]
        additions.append((i, success_text))
    
    if "What new skills have I developed?" in text:
        skills_text = [
            "I developed skills in structuring a larger program into two modules (engine vs. GUI). I used Pygame for the game loop, drawing, and mouse input, and Tkinter for dialogs, which required combining two libraries in one application. I implemented standard chess rules (move generation, check detection, castling, en passant, promotion) and data structures such as the board representation and move log. I practised event-driven design (clicks, key presses, timer updates) and simple file I/O for user credentials. I also improved code readability by adding comments and keeping the engine separate from the display logic, which made debugging and testing easier."
        ]
        additions.append((i, skills_text))

# Add paragraphs - process in reverse order to maintain indices
for para_idx, content in sorted(additions, key=lambda x: x[0], reverse=True):
    para = doc.paragraphs[para_idx]
    # Insert content after this paragraph
    # We insert before the next paragraph, or at end if last
    next_idx = para_idx + 1
    if next_idx < len(doc.paragraphs):
        # Insert before next paragraph (so it appears after current)
        # Insert in reverse order so they appear in correct order
        next_para = doc.paragraphs[next_idx]
        for line in reversed(content):
            new_para = next_para.insert_paragraph_before(line)
            new_para.style = para.style
    else:
        # Last paragraph - add to end in correct order
        for line in content:
            new_para = doc.add_paragraph(line)
            new_para.style = para.style

# Add screenshots after "My test screenshots:"
screenshot_files = [
    ("01_login.png", "Login screen"),
    ("02_board.png", "Board at start"),
    ("04_highlights.png", "Valid moves highlighted"),
    ("05_clocks.png", "Clocks and move history"),
    ("06_checkmate.png", "Checkmate"),
    ("06_promotion.png", "Promotion menu"),
    ("06_restart.png", "Restart menu")
]

# Add images - find the screenshot section and add images there
# We'll add them at the end of the document for simplicity
# (User can manually move them to the correct section if needed)
for img_file, caption in screenshot_files:
    img_path = os.path.join(screenshots_dir, img_file)
    if os.path.exists(img_path):
        try:
            # Add caption
            caption_para = doc.add_paragraph(f"Figure: {caption}")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add image paragraph
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_para.add_run()
            img_run.add_picture(img_path, width=Inches(5))
        except Exception as e:
            print(f"Warning: Could not add image {img_file}: {e}")

# Save the document
output_path = os.path.join(code_dir, "Chess_Game_Project_Report_Final.docx")
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
print(f"File saved to: {output_path}")
